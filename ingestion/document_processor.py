"""
Document processing module using unstructured.io
"""
from pathlib import Path
from typing import List, Tuple, Optional
import os
from unstructured.partition.text import partition_text
from unstructured.documents.elements import Element
import logging

from config.settings import UPLOAD_DIR
from utils.logger import get_logger

logger = get_logger(__name__)

# Disable unstructured API to avoid network calls
os.environ["UNSTRUCTURED_API_KEY"] = ""
os.environ["UNSTRUCTURED_API_URL"] = ""


class DocumentProcessor:
    """Process various document formats using unstructured.io"""
    
    def __init__(self):
        self.logger = logger
    
    def process_document(self, file_path: str) -> Tuple[str, List[str]]:
        """
        Process document and extract text with structure
        
        Args:
            file_path: Path to the document
            
        Returns:
            Tuple of (raw_text, structured_elements)
        """
        try:
            self.logger.info(f"Processing document: {file_path}")
            
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext == ".pdf":
                raw_text = self._process_pdf(file_path)
            elif file_ext in [".docx", ".doc"]:
                raw_text = self._process_docx(file_path)
            elif file_ext == ".txt":
                raw_text = self._process_text(file_path)
            else:
                raw_text = self._process_text(file_path)
            
            # Create simple structured elements (one per paragraph)
            structured_elements = self._structure_text(raw_text)
            
            self.logger.info(f"Successfully processed document")
            return raw_text, structured_elements
            
        except Exception as e:
            self.logger.error(f"Error processing document: {str(e)}")
            raise
    
    def _process_pdf(self, file_path: str) -> str:
        """Process PDF documents using PyPDF2"""
        try:
            self.logger.info("Processing PDF...")
            try:
                import PyPDF2
                text_parts = []
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text = page.extract_text()
                        if text.strip():
                            text_parts.append(text)
                return '\n\n'.join(text_parts)
            except ImportError:
                self.logger.warning("PyPDF2 not available, using pdfplumber...")
                import pdfplumber
                text_parts = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text.strip():
                            text_parts.append(text)
                return '\n\n'.join(text_parts)
        except Exception as e:
            self.logger.error(f"PDF processing failed: {str(e)}")
            raise
    
    def _process_docx(self, file_path: str) -> str:
        """Process DOCX documents using python-docx"""
        try:
            self.logger.info("Processing DOCX...")
            from docx import Document
            doc = Document(file_path)
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    table_rows.append(' | '.join(row_cells))
                if table_rows:
                    text_parts.append('\n'.join(table_rows))
            
            return '\n\n'.join(text_parts) if text_parts else ""
        except Exception as e:
            self.logger.error(f"DOCX processing failed: {str(e)}")
            raise
    
    def _process_text(self, file_path: str) -> str:
        """Process TXT documents"""
        try:
            self.logger.info("Processing TXT...")
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            self.logger.error(f"Text processing failed: {str(e)}")
            raise
    
    def _structure_text(self, text: str) -> List[str]:
        """
        Convert raw text into structured elements (paragraphs)
        """
        # Split by double newlines to get paragraphs
        paragraphs = text.split('\n\n')
        structured = []
        
        for para in paragraphs:
            para = para.strip()
            if para and len(para) > 10:  # Only include meaningful paragraphs
                # Tag with type
                if any(para.startswith(f"{i}.") for i in range(1, 10)):
                    tagged = f"[numbered_list] {para}"
                elif para.startswith(('-', '•', '*')):
                    tagged = f"[bullet] {para}"
                elif len(para) < 100 and para.isupper():
                    tagged = f"[heading] {para}"
                else:
                    tagged = f"[paragraph] {para}"
                structured.append(tagged)
        
        return structured
    
    def extract_metadata(self, structured_elements: List[str]) -> dict:
        """Extract metadata from structured elements"""
        metadata = {
            "total_elements": len(structured_elements),
            "element_types": self._count_element_types(structured_elements),
            "total_text_length": sum(len(elem) for elem in structured_elements),
        }
        return metadata
    
    def _count_element_types(self, structured_elements: List[str]) -> dict:
        """Count different element types from tagged text"""
        type_counts = {}
        for element in structured_elements:
            # Extract type from [type] prefix
            if element.startswith('[') and ']' in element:
                elem_type = element.split(']')[0][1:]
                type_counts[elem_type] = type_counts.get(elem_type, 0) + 1
        return type_counts
